"""
Resume parser — handles PDF (pdfplumber) and DOCX (python-docx).

Design:
- `parse_resume(file_bytes, filename)` is the single public entry point.
- Returns a dict matching the ParsedResume schema.
- Raises ParseError with a descriptive message on failure.
- Never crashes silently — every exception is caught and re-raised as ParseError.
"""
from __future__ import annotations

import io
import re
import logging
from dataclasses import dataclass, field
from typing import Any

logger = logging.getLogger(__name__)

# ── Section heading keywords ───────────────────────────────────────────────

SECTION_PATTERNS: dict[str, list[str]] = {
    "summary":        [r"summary", r"objective", r"profile", r"about me", r"professional summary"],
    "education":      [r"education", r"academic", r"qualification", r"degree"],
    "experience":     [r"experience", r"employment", r"work history", r"career", r"professional background"],
    "projects":       [r"project", r"portfolio", r"personal project", r"side project"],
    "skills":         [r"skill", r"technical skill", r"core competenc", r"expertise", r"technologies"],
    "certifications": [r"certif", r"license", r"accreditation", r"credential"],
    "extras":         [r"publication", r"award", r"honor", r"volunteer", r"language", r"interest", r"activit"],
}

DEGREE_KEYWORDS = {
    "phd": ["ph.d", "phd", "doctorate", "doctoral"],
    "masters": ["master", "m.s.", "msc", "mba", "m.eng", "m.a.", "meng", "m.sc"],
    "bachelors": ["bachelor", "b.s.", "b.a.", "b.eng", "b.tech", "be ", "bs ", "ba ", "bsc", "btech", "b.sc"],
    "associate": ["associate"],
    "diploma": ["diploma", "certificate"],
    "high_school": ["high school", "secondary school", "ged", "hsc", "ssc"],
}


class ParseError(Exception):
    """Raised when a resume file cannot be parsed."""


# ── Public entry point ─────────────────────────────────────────────────────

def parse_resume(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """
    Parse a PDF or DOCX resume from raw bytes.
    Returns a dict matching the ParsedResume schema.
    Raises ParseError if the file cannot be processed.
    """
    if not file_bytes:
        raise ParseError("File is empty.")

    lower = filename.lower()
    try:
        if lower.endswith(".pdf"):
            text, sections = _parse_pdf(file_bytes)
        elif lower.endswith(".docx"):
            text, sections = _parse_docx(file_bytes)
        else:
            raise ParseError(f"Unsupported file type: '{filename}'. Only PDF and DOCX are accepted.")
    except ParseError:
        raise
    except Exception as exc:
        raise ParseError(f"Failed to read file '{filename}': {exc}") from exc

    if not text or len(text.strip()) < 50:
        raise ParseError(f"'{filename}' appears to contain no readable text (possibly image-only PDF or corrupted file).")

    return _extract_structured_data(text, sections)


# ── PDF parsing ────────────────────────────────────────────────────────────

def _parse_pdf(file_bytes: bytes) -> tuple[str, dict[str, str]]:
    """Extract full text and section map from a PDF."""
    try:
        import pdfplumber
    except ImportError:
        raise ParseError("pdfplumber is not installed. Run: pip install pdfplumber")

    pages_text: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        if len(pdf.pages) == 0:
            raise ParseError("PDF has no pages.")
        for page in pdf.pages:
            page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
            if page_text:
                pages_text.append(page_text)

    full_text = "\n".join(pages_text)
    sections = _split_into_sections(full_text)
    return full_text, sections


# ── DOCX parsing ───────────────────────────────────────────────────────────

def _parse_docx(file_bytes: bytes) -> tuple[str, dict[str, str]]:
    """Extract full text and section map from a DOCX."""
    try:
        from docx import Document
    except ImportError:
        raise ParseError("python-docx is not installed. Run: pip install python-docx")

    doc = Document(io.BytesIO(file_bytes))
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Also extract table text (some resumes use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in lines:
                    lines.append(cell_text)

    full_text = "\n".join(lines)
    sections = _split_into_sections(full_text)
    return full_text, sections


# ── Section splitting ──────────────────────────────────────────────────────

def _split_into_sections(text: str) -> dict[str, str]:
    """
    Split raw resume text into named sections.
    Returns dict mapping section_name → section_text.
    Unrecognised leading content goes into "header".
    """
    lines = text.split("\n")
    sections: dict[str, list[str]] = {"header": []}
    current_section = "header"

    for line in lines:
        stripped = line.strip()
        if not stripped:
            sections.setdefault(current_section, []).append("")
            continue

        detected = _detect_section_heading(stripped)
        if detected and len(stripped) < 60:  # headings are short
            current_section = detected
            sections.setdefault(current_section, [])
        else:
            sections.setdefault(current_section, []).append(stripped)

    return {k: "\n".join(v).strip() for k, v in sections.items()}


def _detect_section_heading(line: str) -> str | None:
    """Return the canonical section name if the line looks like a section heading."""
    normalized = line.lower().strip(" :-–—\t")
    for section_name, patterns in SECTION_PATTERNS.items():
        for pattern in patterns:
            if re.search(pattern, normalized):
                return section_name
    return None


# ── Structured extraction ──────────────────────────────────────────────────

def _extract_structured_data(full_text: str, sections: dict[str, str]) -> dict[str, Any]:
    """Combine header + section parsing into the ParsedResume dict shape."""
    header_text = sections.get("header", "") + "\n" + full_text[:500]

    return {
        "name": _extract_name(header_text),
        "email": _extract_email(full_text),
        "phone": _extract_phone(full_text),
        "location": _extract_location(header_text),
        "linkedin": _extract_url(full_text, r"linkedin\.com/in/[\w\-]+"),
        "github": _extract_url(full_text, r"github\.com/[\w\-]+"),
        "website": _extract_url(full_text, r"https?://(?!linkedin|github)[\w\-\.]+\.\w+"),
        "summary": sections.get("summary", "").strip(),
        "education": _parse_education_section(sections.get("education", "")),
        "experience": _parse_experience_section(sections.get("experience", "")),
        "projects": _parse_projects_section(sections.get("projects", "")),
        "skills": _parse_skills_section(sections.get("skills", "")),
        "certifications": _parse_certifications_section(sections.get("certifications", "")),
        "extras": _parse_extras_section(sections.get("extras", "")),
    }


# ── Contact extraction helpers ─────────────────────────────────────────────

def _extract_email(text: str) -> str:
    m = re.search(r"[\w.+-]+@[\w.-]+\.[a-zA-Z]{2,}", text)
    return m.group(0).lower() if m else ""


def _extract_phone(text: str) -> str:
    m = re.search(r"[\+\(]?[\d\s\-\.\(\)]{7,15}", text)
    if m:
        candidate = re.sub(r"\s+", "", m.group(0))
        if len(candidate) >= 7:
            return candidate
    return ""


def _extract_url(text: str, pattern: str) -> str:
    m = re.search(pattern, text, re.IGNORECASE)
    return m.group(0) if m else ""


def _extract_name(header: str) -> str:
    """Heuristic: first non-empty line that looks like a name (no @ or digits)."""
    for line in header.split("\n"):
        line = line.strip()
        if (line and 3 < len(line) < 60
                and "@" not in line
                and not re.search(r"\d{3,}", line)
                and not re.search(r"(http|www|linkedin|github|resume|cv)\b", line, re.I)):
            return line
    return ""


def _extract_location(text: str) -> str:
    """Look for City, State/Country patterns."""
    m = re.search(
        r"\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)*),\s*([A-Z]{2}|[A-Z][a-z]+(?:\s[A-Z][a-z]+)*)\b",
        text
    )
    return m.group(0) if m else ""


# ── Section parsers ────────────────────────────────────────────────────────

def _parse_education_section(text: str) -> list[dict]:
    """Extract education entries from the education section."""
    if not text.strip():
        return []

    entries: list[dict] = []
    # Split by blank lines or lines that start a new degree entry
    blocks = re.split(r"\n\s*\n", text)
    if len(blocks) == 1:
        # Try splitting by degree keywords
        blocks = re.split(r"(?=\b(?:bachelor|master|phd|associate|diploma|b\.s|m\.s|m\.b\.a|b\.tech)\b)", text, flags=re.I)

    for block in blocks:
        block = block.strip()
        if not block:
            continue
        entry = _parse_single_education(block)
        if entry["degree"] or entry["institution"]:
            entries.append(entry)

    return entries if entries else []


def _parse_single_education(text: str) -> dict:
    degree_level = ""
    degree_text = ""
    for level, keywords in DEGREE_KEYWORDS.items():
        for kw in keywords:
            if kw.lower() in text.lower():
                degree_level = level
                degree_text = _find_context(text, kw, 60)
                break
        if degree_level:
            break

    # Year extraction
    years = re.findall(r"\b(19|20)\d{2}\b", text)
    year_start = int(years[0]) if len(years) >= 2 else None
    year_end = int(years[-1]) if years else None

    # GPA
    gpa_m = re.search(r"(?:gpa|grade)[:\s]+(\d\.\d+)", text, re.I)
    gpa = float(gpa_m.group(1)) if gpa_m else None

    # Field of study — after "in" or "of"
    field_m = re.search(r"\b(?:in|of)\s+([A-Za-z\s&]+?)(?:\s*[,\n\(]|$)", degree_text or text)
    field = field_m.group(1).strip() if field_m else ""

    # Institution — look for University/College/Institute
    inst_m = re.search(r"([A-Z][A-Za-z\s&']+(?:University|College|Institute|School|Academy))", text)
    institution = inst_m.group(1).strip() if inst_m else ""

    return {
        "degree": degree_text or degree_level,
        "field": field,
        "institution": institution,
        "year_start": year_start,
        "year_end": year_end,
        "gpa": gpa,
    }


def _parse_experience_section(text: str) -> list[dict]:
    """Extract work experience entries."""
    if not text.strip():
        return []

    # Split by company/title blocks (blank line separation usually works)
    blocks = re.split(r"\n\s*\n", text)
    entries = []
    for block in blocks:
        block = block.strip()
        if not block or len(block) < 20:
            continue
        entry = _parse_single_experience(block)
        if entry["title"] or entry["company"]:
            entries.append(entry)
    return entries


def _parse_single_experience(text: str) -> dict:
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    # First line is often "Title | Company" or "Title at Company"
    title = ""
    company = ""
    if lines:
        first = lines[0]
        if " | " in first:
            parts = first.split(" | ", 1)
            title, company = parts[0].strip(), parts[1].strip()
        elif " at " in first.lower():
            parts = re.split(r"\s+at\s+", first, flags=re.I, maxsplit=1)
            title, company = parts[0].strip(), parts[1].strip()
        elif " - " in first:
            parts = first.split(" - ", 1)
            title, company = parts[0].strip(), parts[1].strip()
        else:
            title = first
            company = lines[1] if len(lines) > 1 else ""

    # Date range
    date_pattern = r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|\d{4})"
    dates = re.findall(date_pattern, " ".join(lines), re.I)

    start_date = dates[0] if len(dates) >= 1 else ""
    end_date = dates[1] if len(dates) >= 2 else ""
    is_current = bool(re.search(r"\bpresent\b|\bcurrent\b|\bnow\b", text, re.I))
    if is_current:
        end_date = "Present"

    # Duration calculation
    duration_months = _calculate_duration(start_date, end_date)

    # Description — everything after the first two lines
    description = "\n".join(lines[2:]) if len(lines) > 2 else "\n".join(lines[1:])

    # Skills mentioned in description
    skills_mentioned = _extract_tech_terms(description)

    return {
        "title": title,
        "company": company,
        "start_date": start_date,
        "end_date": end_date,
        "is_current": is_current,
        "duration_months": duration_months,
        "description": description,
        "skills_mentioned": skills_mentioned,
    }


def _parse_projects_section(text: str) -> list[dict]:
    """Extract project entries."""
    if not text.strip():
        return []

    blocks = re.split(r"\n\s*\n", text)
    entries = []
    for block in blocks:
        block = block.strip()
        if not block or len(block) < 15:
            continue
        lines = [l.strip() for l in block.split("\n") if l.strip()]
        title = lines[0] if lines else ""
        description = "\n".join(lines[1:]) if len(lines) > 1 else ""
        links = re.findall(r"https?://\S+|github\.com/\S+", block)
        skills = _extract_tech_terms(block)
        entries.append({
            "title": title,
            "description": description,
            "skills": skills,
            "links": links,
        })
    return entries


def _parse_skills_section(text: str) -> list[str]:
    """Extract skills list from the skills section."""
    if not text.strip():
        return []

    # Remove common label prefixes like "Languages:", "Frameworks:"
    text = re.sub(r"^[A-Za-z\s/]+:\s*", "", text, flags=re.MULTILINE)

    # Split on commas, pipes, bullets, semicolons, newlines
    raw = re.split(r"[,|•·\n;/]+", text)
    skills = []
    for item in raw:
        item = item.strip(" .-–—\t")
        if item and 1 < len(item) < 50 and not item.lower().startswith("skill"):
            skills.append(item)
    return list(dict.fromkeys(skills))  # deduplicate, preserve order


def _parse_certifications_section(text: str) -> list[str]:
    """Extract certification names."""
    if not text.strip():
        return []
    lines = [l.strip(" .-–•·\t") for l in text.split("\n") if l.strip()]
    return [l for l in lines if l and len(l) < 150]


def _parse_extras_section(text: str) -> dict[str, Any]:
    """Parse publications, awards, languages, volunteer work."""
    result: dict[str, list] = {
        "publications": [],
        "awards": [],
        "languages": [],
        "volunteer": [],
        "other": [],
    }
    if not text.strip():
        return result

    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines:
        lower = line.lower()
        if any(w in lower for w in ["publish", "paper", "journal", "conference", "arxiv"]):
            result["publications"].append(line)
        elif any(w in lower for w in ["award", "honor", "prize", "winner", "scholarship"]):
            result["awards"].append(line)
        elif any(w in lower for w in ["english", "spanish", "french", "german", "mandarin", "hindi", "language"]):
            result["languages"].append(line)
        elif any(w in lower for w in ["volunteer", "nonprofit", "charity", "community"]):
            result["volunteer"].append(line)
        else:
            result["other"].append(line)
    return result


# ── Utility helpers ────────────────────────────────────────────────────────

def _find_context(text: str, keyword: str, window: int = 60) -> str:
    """Return the substring around a keyword match."""
    idx = text.lower().find(keyword.lower())
    if idx == -1:
        return ""
    start = max(0, idx - 5)
    end = min(len(text), idx + window)
    return text[start:end].strip()


_TECH_TERMS = re.compile(
    r"\b(python|javascript|typescript|java|c\+\+|c#|go|rust|ruby|swift|kotlin|scala|r\b|"
    r"react|angular|vue|node|express|django|flask|fastapi|spring|rails|laravel|"
    r"tensorflow|pytorch|keras|sklearn|scikit|numpy|pandas|matplotlib|"
    r"aws|gcp|azure|docker|kubernetes|k8s|terraform|ansible|jenkins|ci/cd|"
    r"sql|postgresql|postgres|mysql|mongodb|redis|elasticsearch|kafka|"
    r"git|linux|bash|html|css|rest|graphql|grpc|microservice|"
    r"machine learning|deep learning|nlp|computer vision|data science|"
    r"agile|scrum|jira|figma|tableau|power bi)\b",
    re.IGNORECASE,
)


def _extract_tech_terms(text: str) -> list[str]:
    """Extract known tech keywords from a blob of text."""
    found = _TECH_TERMS.findall(text)
    return list(dict.fromkeys(t.strip() for t in found))


def _calculate_duration(start: str, end: str) -> int | None:
    """Return approximate duration in months between two date strings."""
    try:
        from dateutil import parser as dparser
        from datetime import datetime

        start_dt = dparser.parse(start, default=datetime(2000, 1, 1))
        if end and end.lower() not in ("present", "current", "now", ""):
            end_dt = dparser.parse(end, default=datetime(2000, 1, 1))
        else:
            end_dt = datetime.now()

        months = (end_dt.year - start_dt.year) * 12 + (end_dt.month - start_dt.month)
        return max(0, months)
    except Exception:
        return None
