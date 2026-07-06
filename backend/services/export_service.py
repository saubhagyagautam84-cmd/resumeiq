"""
ATS-safe resume export — generates PDF (reportlab) and DOCX (python-docx).

Design rules for ATS safety:
- Single column layout only.
- No text boxes, tables, graphics, or images.
- Standard fonts and headings.
- Machine-readable text throughout.
- After generation, the file is passed back through the resume parser as a self-test.
"""
from __future__ import annotations

import io
import logging
from typing import Any

logger = logging.getLogger(__name__)


# ── PDF export (reportlab) ─────────────────────────────────────────────────

def export_pdf(resume_data: dict[str, Any]) -> bytes:
    """Generate an ATS-safe PDF resume. Returns raw bytes."""
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
    except ImportError:
        raise RuntimeError("reportlab is not installed. Run: pip install reportlab")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    name_style = ParagraphStyle("Name", parent=styles["Heading1"], fontSize=18, spaceAfter=4, alignment=TA_CENTER)
    contact_style = ParagraphStyle("Contact", parent=styles["Normal"], fontSize=9, spaceAfter=2, alignment=TA_CENTER)
    section_style = ParagraphStyle("Section", parent=styles["Heading2"], fontSize=11, spaceBefore=10, spaceAfter=2,
                                   textColor=colors.HexColor("#1a1a1a"), borderPad=0)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, spaceAfter=2, leading=13)
    bullet_style = ParagraphStyle("Bullet", parent=styles["Normal"], fontSize=10, leftIndent=12, spaceAfter=1, leading=13)
    small_style = ParagraphStyle("Small", parent=styles["Normal"], fontSize=9, spaceAfter=1, textColor=colors.HexColor("#555555"))

    story = []

    # Header
    name = resume_data.get("name") or "Your Name"
    story.append(Paragraph(name, name_style))

    contact_parts = [p for p in [
        resume_data.get("email"),
        resume_data.get("phone"),
        resume_data.get("location"),
        resume_data.get("linkedin"),
        resume_data.get("github"),
        resume_data.get("website"),
    ] if p]
    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), contact_style))

    story.append(Spacer(1, 4))

    def section_hr():
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc")))

    # Summary
    summary = resume_data.get("summary", "").strip()
    if summary:
        story.append(Paragraph("SUMMARY", section_style))
        section_hr()
        story.append(Paragraph(summary, body_style))
        story.append(Spacer(1, 6))

    # Experience
    experience = resume_data.get("experience", [])
    if experience:
        story.append(Paragraph("EXPERIENCE", section_style))
        section_hr()
        for exp in experience:
            title = exp.get("title", "")
            company = exp.get("company", "")
            start = exp.get("start_date", "")
            end = exp.get("end_date", "") or ("Present" if exp.get("is_current") else "")
            date_str = f"{start} – {end}".strip(" –")

            story.append(Paragraph(f"<b>{title}</b> — {company}", body_style))
            if date_str:
                story.append(Paragraph(date_str, small_style))
            desc = exp.get("description", "")
            if desc:
                for line in desc.split("\n"):
                    line = line.strip(" •·-")
                    if line:
                        story.append(Paragraph(f"• {line}", bullet_style))
            story.append(Spacer(1, 4))

    # Education
    education = resume_data.get("education", [])
    if education:
        story.append(Paragraph("EDUCATION", section_style))
        section_hr()
        for edu in education:
            degree = edu.get("degree", "")
            field = edu.get("field", "")
            institution = edu.get("institution", "")
            year_end = edu.get("year_end", "")
            gpa = edu.get("gpa")

            degree_line = " in ".join(filter(None, [degree, field]))
            story.append(Paragraph(f"<b>{degree_line}</b>", body_style))
            detail_parts = [institution, str(year_end) if year_end else ""]
            if gpa:
                detail_parts.append(f"GPA: {gpa}")
            story.append(Paragraph(" | ".join(p for p in detail_parts if p), small_style))
            story.append(Spacer(1, 4))

    # Projects
    projects = resume_data.get("projects", [])
    if projects:
        story.append(Paragraph("PROJECTS", section_style))
        section_hr()
        for proj in projects:
            title = proj.get("title", "")
            desc = proj.get("description", "")
            skills = proj.get("skills", [])
            links = proj.get("links", [])

            story.append(Paragraph(f"<b>{title}</b>", body_style))
            if desc:
                story.append(Paragraph(desc, bullet_style))
            if skills:
                story.append(Paragraph(f"Tech: {', '.join(skills)}", small_style))
            if links:
                story.append(Paragraph(f"Link: {links[0]}", small_style))
            story.append(Spacer(1, 4))

    # Skills
    skills = resume_data.get("skills", [])
    if skills:
        story.append(Paragraph("SKILLS", section_style))
        section_hr()
        story.append(Paragraph(", ".join(skills), body_style))
        story.append(Spacer(1, 6))

    # Certifications
    certs = resume_data.get("certifications", [])
    if certs:
        story.append(Paragraph("CERTIFICATIONS", section_style))
        section_hr()
        for cert in certs:
            story.append(Paragraph(f"• {cert}", bullet_style))
        story.append(Spacer(1, 6))

    doc.build(story)
    return buf.getvalue()


# ── DOCX export (python-docx) ──────────────────────────────────────────────

def export_docx(resume_data: dict[str, Any]) -> bytes:
    """Generate an ATS-safe DOCX resume. Returns raw bytes."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = Document()

    # Remove default margins (set to 1 inch)
    section = doc.sections[0]
    from docx.shared import Inches
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    def add_heading(text: str, level: int = 1):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def add_para(text: str, bold: bool = False, italic: bool = False, small: bool = False) -> Any:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(9 if small else 10)
        return p

    def add_bullet(text: str):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text).font.size = Pt(10)
        return p

    # Name
    name_para = doc.add_heading(resume_data.get("name") or "Your Name", level=1)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact line
    contact_parts = [p for p in [
        resume_data.get("email"),
        resume_data.get("phone"),
        resume_data.get("location"),
        resume_data.get("linkedin"),
        resume_data.get("github"),
    ] if p]
    if contact_parts:
        cp = doc.add_paragraph(" | ".join(contact_parts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.font.size = Pt(9)

    # Summary
    summary = resume_data.get("summary", "").strip()
    if summary:
        add_heading("SUMMARY", level=2)
        add_para(summary)

    # Experience
    experience = resume_data.get("experience", [])
    if experience:
        add_heading("EXPERIENCE", level=2)
        for exp in experience:
            title = exp.get("title", "")
            company = exp.get("company", "")
            start = exp.get("start_date", "")
            end = exp.get("end_date", "") or ("Present" if exp.get("is_current") else "")
            date_str = f"{start} – {end}".strip(" –")
            add_para(f"{title} — {company}", bold=True)
            if date_str:
                add_para(date_str, small=True)
            desc = exp.get("description", "")
            if desc:
                for line in desc.split("\n"):
                    line = line.strip(" •·-")
                    if line:
                        add_bullet(line)

    # Education
    education = resume_data.get("education", [])
    if education:
        add_heading("EDUCATION", level=2)
        for edu in education:
            degree = edu.get("degree", "")
            field = edu.get("field", "")
            institution = edu.get("institution", "")
            year_end = edu.get("year_end", "")
            degree_line = " in ".join(filter(None, [degree, field]))
            add_para(degree_line, bold=True)
            detail = " | ".join(p for p in [institution, str(year_end) if year_end else ""] if p)
            if detail:
                add_para(detail, small=True)

    # Projects
    projects = resume_data.get("projects", [])
    if projects:
        add_heading("PROJECTS", level=2)
        for proj in projects:
            add_para(proj.get("title", ""), bold=True)
            if proj.get("description"):
                add_bullet(proj["description"])
            if proj.get("skills"):
                add_para(f"Tech: {', '.join(proj['skills'])}", small=True)
            if proj.get("links"):
                add_para(f"Link: {proj['links'][0]}", small=True)

    # Skills
    skills = resume_data.get("skills", [])
    if skills:
        add_heading("SKILLS", level=2)
        add_para(", ".join(skills))

    # Certifications
    certs = resume_data.get("certifications", [])
    if certs:
        add_heading("CERTIFICATIONS", level=2)
        for cert in certs:
            add_bullet(cert)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Self-test ──────────────────────────────────────────────────────────────

def self_test_export(file_bytes: bytes, filename: str) -> list[str]:
    """
    Run the generated resume back through the parser.
    Returns a list of warning strings (empty if the resume parses cleanly).
    """
    from services.resume_parser import parse_resume, ParseError
    warnings: list[str] = []
    try:
        parsed = parse_resume(file_bytes, filename)
        if not parsed.get("name"):
            warnings.append("Name could not be extracted from the exported file — check formatting.")
        if not parsed.get("experience") and not parsed.get("education"):
            warnings.append("Experience and education sections were not detected in the exported file — ATS may struggle.")
        if not parsed.get("skills"):
            warnings.append("Skills section was not detected in the exported file.")
    except ParseError as exc:
        warnings.append(f"Self-test failed: the exported resume could not be re-parsed ({exc}). Check for formatting issues.")
    return warnings
